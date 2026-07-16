import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    """Ajoute une couleur de fond à une cellule de tableau"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def creer_rapport_complet():
    doc = docx.Document()
    
    # Couleurs professionnelles
    COLOR_PRIMARY = RGBColor(26, 54, 93)     # Bleu Marine
    COLOR_SECONDARY = RGBColor(43, 108, 176)  # Bleu Acier
    COLOR_TEXT = RGBColor(45, 55, 72)        # Gris foncé
    
    # Configuration de la police par défaut
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = COLOR_TEXT

    # --- TITRE PRINCIPAL ---
    title = doc.add_paragraph()
    title_run = title.add_run("RAPPORT TECHNIQUE DE FIN DE PHASE 1")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_PRIMARY
    
    # --- METADONNÉES ---
    meta = doc.add_paragraph()
    meta.add_run("Étudiante : ").bold = True
    meta.add_run("Mounia Zouak  |  ")
    meta.add_run("Filière : ").bold = True
    meta.add_run("Digital Data & AI / Big Data & AI\n")
    meta.add_run("Dépôt de code officiel : ").bold = True
    meta.add_run("https://github.com/mouniazouak/phase1-jira-test").font.color.rgb = COLOR_SECONDARY
    
    doc.add_paragraph("-" * 80)

    # --- SECTION 1 ---
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Objectifs Rappelés (Phase 1)")
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = COLOR_PRIMARY
    
    objectifs = [
        "Cartographier les sources de données.",
        "Définir le modèle de carte projet cible.",
        "Réaliser un POC (Proof of Concept) de mise à jour automatique."
    ]
    for obj in objectifs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(obj)

    # --- SECTION 2 ---
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Réalisations Techniques & Cartographie Applicative")
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph(
        "L'environnement de développement a été isolé, configuré et interconnecté avec les dépendances "
        "clés nécessaires à l'ensemble du cycle de vie du projet. L'architecture logicielle de la Phase 1 s'appuie sur :"
    )
    
    techs = [
        ("Ingestion & Modélisation (Jira) : ", "Intégration de la bibliothèque jira permettant d'interroger l'API, d'extraire les structures de tickets (Epics, Sprints, Stories) et de poser les bases du modèle de carte projet cible."),
        ("Pipeline & Exposition API (FastAPI & Uvicorn) : ", "Déploiement d'une architecture backend asynchrone pour piloter la logique de mise à jour automatique et l'écoute de futurs webhooks de synchronisation."),
        ("Traitement Analytique : ", "Exploitation des packages pandas et numpy pour la structuration, le nettoyage et le reformatage des données brutes récoltées."),
        ("Persistance (Base de données) : ", "Utilisation du connecteur pymongo pour historiser les états de projets dans une base NoSQL MongoDB."),
        ("Moteurs de Rendu pour la Phase 2 : ", "Pré-intégration des modules de génération automatique de livrables (tableaux Excel stylisés avec xlsxwriter, rapports PDF avec weasyprint/reportlab et présentations avec python-pptx).")
    ]
    for titre, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(titre)
        r.bold = True
        r.font.color.rgb = COLOR_SECONDARY
        p.add_run(desc)

    # --- SECTION 3 (SCHÉMA) ---
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Schéma Architectural du Système")
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = COLOR_PRIMARY
    
    schema_text = (
        "       ┌────────────────────────────────────────────────────────┐\n"
        "       │                      Source Externe                    │\n"
        "       │                         [ JIRA ]                       │\n"
        "       └────────────────────────────┬───────────────────────────?\n"
        "                                    │\n"
        "                                    │ (Extraction de données via API)\n"
        "                                    ▼\n"
        "       ┌────────────────────────────────────────────────────────┐\n"
        "       │                   C:\\phase1-jira-test                  │\n"
        "       │                                                        │\n"
        "       │  ┌──────────────────────────────────────────────────┐  │\n"
        "       │  │                1. Ingestion & API                │  │\n"
        "       │  │               [ jira ] [ fastapi ]               │  │\n"
        "       │  └─────────────────────────┬────────────────────────┘  │\n"
        "       │                            │                           │\n"
        "       │                            ▼                           │\n"
        "       │  ┌──────────────────────────────────────────────────┐  │\n"
        "       │  │             2. Traitement Analytique             │  │\n"
        "       │  │               [ pandas ] [ numpy ]               │  │\n"
        "       │  └─────────────────────────┬────────────────────────┘  │\n"
        "       │                            │                           │\n"
        "       │                            ▼                           │\n"
        "       │  ┌──────────────────────────────────────────────────┐  │\n"
        "       │  │             3. Persistance & Stockage            │  │\n"
        "       │  │                    [ pymongo ]                   │  │\n"
        "       │  └─────────────────────────┬────────────────────────┘  │\n"
        "       │                            │                           │\n"
        "       │                            ▼                           │\n"
        "       │  ┌──────────────────────────────────────────────────┐  │\n"
        "       │  │            4. Moteurs d'Export (Visual)          │  │\n"
        "       │  │ [ weasyprint ] [ reportlab ] [ pptx ] [ xlsx ]   │  │\n"
        "       │  └──────────────────────────────────────────────────┘  │\n"
        "       │                                                        │\n"
        "       └────────────────────────────────────────────────────────┘"
    )
    
    p_schema = doc.add_paragraph()
    p_schema_run = p_schema.add_run(schema_text)
    p_schema_run.font.name = 'Consolas'
    p_schema_run.font.size = Pt(9.5)
    p_schema_run.font.color.rgb = RGBColor(0, 0, 0)

    # --- SECTION 4 (TABLEAU DÉTAILLÉ) ---
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Détail d'Utilisation des Composants par Bloc")
    h4_run.font.size = Pt(14)
    h4_run.font.bold = True
    h4_run.font.color.rgb = COLOR_PRIMARY
    
    table_data = [
        ("Bloc Architectural", "Bibliothèques / Outils", "Rôle et Usage Précis"),
        ("1. Ingestion & API", "jira (Python SDK)\nfastapi & uvicorn\npydantic", "Connexion sécurisée aux instances Jira, extraction des Epics/Sprints. Exposition d'endpoints de synchronisation et écoute de Webhooks."),
        ("2. Traitement Analytique", "pandas\nnumpy", "Modélisation des données brutes Jira sous forme de DataFrames. Opérations de tri, nettoyage et calculs de dérives."),
        ("3. Persistance & Stockage", "pymongo (MongoDB)", "Sauvegarde et historisation souple des états sous forme de documents NoSQL JSON polymorphes."),
        ("4. Moteurs d'Export", "weasyprint & reportlab\npython-pptx\nxlsxwriter", "Pré-configuration des moteurs de génération de rapports PDF automatisés, de diaporamas PPTX et de classeurs de données Excel."),
        ("5. Qualité & Environnement", "pytest\npython-dotenv\ngit & .gitignore", "Validation des règles métiers (dossier /tests), isolation des clés d'API et variables sensibles, contrôle de version de l'arborescence.")
    ]
    
    table = doc.add_table(rows=len(table_data), cols=3)
    table.style = 'Light Shading Accent 1'
    
    for i, row_data in enumerate(table_data):
        row = table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            # Formatage de la ligne d'entête
            if i == 0:
                set_cell_background(cell, "1A365D")
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph().paragraph_format.space_before = Pt(10)

    # --- SECTION 5 ---
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. Conclusion et Passage à la Phase 2")
    h5_run.font.size = Pt(14)
    h5_run.font.bold = True
    h5_run.font.color.rgb = COLOR_PRIMARY
    
    doc.add_paragraph(
        "Le POC d'initialisation et d'ingestion de flux automatique est validé. Les fondations applicatives étant "
        "stables, packagées et archivées de façon sécurisée, le projet bascule dès à présent sur la Phase 2, axée sur :"
    )
    
    suivants = [
        "La génération et l'export automatique des reportings visuels (PDF, Excel, PPTX).",
        "L'implémentation des modèles de calculs de dérive planning et d'alertes.",
        "L'interfaçage progressif avec le protocole Microsoft Teams."
    ]
    for suiv in suivants:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(suiv)

    # Sauvegarde finale
    filename = "Rapport_Architecture_Complet_Phase1.docx"
    doc.save(filename)
    print(f"Succès : Le fichier '{filename}' a été généré avec succès.")

if __name__ == "__main__":
    creer_rapport_complet()